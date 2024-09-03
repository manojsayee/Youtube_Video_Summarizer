from pytube import YouTube
from huggingsound import SpeechRecognitionModel
import torch
import librosa
import soundfile as sf
from transformers import pipeline
import os
from fastapi.templating import Jinja2Templates
import docx
from fastapi.responses import FileResponse
from pathlib import Path


templates = Jinja2Templates(directory="templates/html")

def summarize(VIDEO_URL):
  # VIDEO_URL = "https://www.youtube.com/watch?v=hWLf6JFbZoo"

  yt = YouTube(VIDEO_URL)

  yt.streams \
    .filter(only_audio = True, file_extension = 'mp4') \
    .first() \
    .download(filename = 'ytaudio.mp4')

  device = "cuda" if torch.cuda.is_available() else "cpu"
  print(device)

  model = SpeechRecognitionModel("jonatasgrosman/wav2vec2-large-xlsr-53-english", device = device)

  os.system('ffmpeg -i ytaudio.mp4 -acodec pcm_s16le -ar 16000 ytaudio.wav')

  input_file = './ytaudio.wav'

  # print(librosa.get_samplerate(input_file))

  stream = librosa.stream(
      input_file,
      block_length=30,
      frame_length=16000,
      hop_length=16000
  )
  print(stream)
  for i,speech in enumerate(stream):
    sf.write(f'content/{i}.wav', speech, 16000)

  audio_path =[]
  for a in range(i+1):
    audio_path.append(f'C:/Users/Manoj/Downloads/project/project/content/{a}.wav') 

  transcriptions = model.transcribe(audio_path)

  full_transcript = ' '

  for item in transcriptions:
    full_transcript += ''.join(item['transcription'])
  
  print(full_transcript)
  summarization = pipeline('summarization')

  summarized_text = summarization(full_transcript)
  print(summarized_text)

  return {'actual_text': full_transcript ,'summarized_text':summarized_text}

def download_as_text(act_text, sum_text)->None:
  try:
    doc = docx.Document()

    doc.add_heading('Actual Text', 0)
    act_para = doc.add_paragraph(act_text)

    doc.add_heading('Summarized Text', 0)
    sum_para = doc.add_paragraph(sum_text)

    # file_path = os.getcwd() + "/" + "summarized_data.docx"

    downloads_path = str(Path.home() / "Downloads")

    file_path = downloads_path + "/" +"summarized_data.docx"
    doc.save(file_path)

    return FileResponse(path=file_path, media_type='text/mp4', filename= "summarized_data.docx")
  except Exception as e:
    print(e)

def delete_files_in_directory(directory_path):
   try:
     files = os.listdir(directory_path)
     for file in files:
       file_path = os.path.join(directory_path, file)
       if os.path.isfile(file_path):
         os.remove(file_path)
     print("All files deleted successfully.")
   except OSError:
     print("Error occurred while deleting files.")